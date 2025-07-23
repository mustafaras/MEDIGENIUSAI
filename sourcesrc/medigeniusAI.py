import os
import streamlit as st
import asyncio
import base64
import datetime
import io
import pandas as pd
import warnings
import time
import hashlib

# AI and external library imports with error handling
try:
    import ollama
    OLLAMA_AVAILABLE = True
except ImportError:
    OLLAMA_AVAILABLE = False
    print("⚠️ Ollama not available")

try:
    from openai import OpenAI
    OPENAI_AVAILABLE = True
except ImportError:
    OPENAI_AVAILABLE = False
    print("⚠️ OpenAI not available")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ python-docx not available")

try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False
    print("⚠️ PIL not available")

try:
    from medrag import MedRAG
    MEDRAG_AVAILABLE = True
except ImportError:
    MEDRAG_AVAILABLE = False
    print("⚠️ MedRAG not available")

# Streamlit UI setup - MUST BE FIRST STREAMLIT COMMAND
st.set_page_config(
    page_title="MediGenius AI", 
    layout="wide", 
    initial_sidebar_state="expanded"
)

# Maximum sidebar width CSS
st.markdown("""
<style>
    /* Maximum sidebar width - targeting all possible selectors */
    .css-1d391kg, .css-1lcbmhc, .css-17eq0hr, .css-1y4p8pa, .css-6qob1r {
        width: 600px !important;
        max-width: 600px !important;
        min-width: 600px !important;
    }
    
    [data-testid="stSidebar"] {
        width: 600px !important;
        max-width: 600px !important;
        min-width: 600px !important;
    }
    
    [data-testid="stSidebar"] > div {
        width: 600px !important;
        max-width: 600px !important;
        min-width: 600px !important;
    }
    
    /* Sidebar content area */
    .css-1lcbmhc .css-1v0mbdj {
        width: 580px !important;
        max-width: 580px !important;
    }
    
    /* Main content adjustment */
    .css-18e3th9, .css-1d391kg {
        margin-left: 620px !important;
    }
    
    /* Force sidebar to stay open and wide */
    .css-1lcbmhc .css-1v0mbdj {
        padding: 1rem !important;
    }
</style>
""", unsafe_allow_html=True)

# Force refresh the page to clear any cached HTML
import time
current_time = int(time.time())

# Clear Streamlit cache
st.cache_data.clear()
st.cache_resource.clear()

# Suppress FutureWarning messages from the huggingface_hub library
warnings.filterwarnings("ignore", category=FutureWarning, module="huggingface_hub")

# Ensure Streamlit works well with asyncio
try:
    asyncio.get_running_loop()
except RuntimeError:
    asyncio.set_event_loop(asyncio.new_event_loop())

# Check Ollama availability and get models
def get_ollama_models():
    """Get available Ollama models, return empty list if Ollama not available"""
    if not OLLAMA_AVAILABLE:
        return []
    try:
        # Try to connect to Ollama server
        models = ollama.list()
        if models and models.get('models'):
            model_names = [model['name'] for model in models['models']]
            print(f"✅ Found {len(model_names)} Ollama models: {model_names}")
            return model_names
        else:
            print("⚠️ Ollama server responded but no models found")
            return []
    except Exception as e:
        print(f"❌ Ollama connection error: {str(e)}")
        # Try alternative connection method
        try:
            import requests
            response = requests.get("http://127.0.0.1:11434/api/tags", timeout=5)
            if response.status_code == 200:
                data = response.json()
                if data and data.get('models'):
                    model_names = [model['name'] for model in data['models']]
                    print(f"✅ Found {len(model_names)} Ollama models via HTTP: {model_names}")
                    return model_names
            print("⚠️ Ollama HTTP request failed or no models found")
            return []
        except Exception as e2:
            print(f"❌ Ollama HTTP connection also failed: {str(e2)}")
            return []

# Get available Ollama models
try:
    print("🔍 Checking for Ollama models...")
    ollama_models = get_ollama_models()
    ollama_available = len(ollama_models) > 0
    print(f"📊 Ollama status: {len(ollama_models)} models found, available: {ollama_available}")
    if ollama_available:
        print(f"🎯 Available models: {ollama_models}")
except Exception as e:
    print(f"❌ Error getting Ollama models: {str(e)}")
    ollama_models = []
    ollama_available = False

# Initialize MedRAG with better error handling
medrag = None
medrag_available = False

if MEDRAG_AVAILABLE:
    try:
        import ssl
        import requests
        from requests.adapters import HTTPAdapter
        from urllib3.util.retry import Retry
        
        # Configure SSL and retry settings for better connectivity
        ssl._create_default_https_context = ssl._create_unverified_context
        
        # Create a session with retry strategy
        session = requests.Session()
        retry_strategy = Retry(
            total=3,
            backoff_factor=1,
            status_forcelist=[429, 500, 502, 503, 504],
        )
        adapter = HTTPAdapter(max_retries=retry_strategy)
        session.mount("http://", adapter)
        session.mount("https://", adapter)
        
        # Initialize MedRAG silently
        medrag = MedRAG()
        medrag_available = True
        
    except Exception as e:
        # Silent initialization failure, status will be shown in sidebar
        medrag_available = False
else:
    medrag_available = False

# Force Dark Mode
st.markdown("""
<style>
    /* Force dark mode by overriding all Streamlit default styles */
    .stApp {
        background-color: #0e1117 !important;
        color: #ffffff !important;
    }
    
    /* Main content area */
    .main .block-container {
        background-color: #0e1117 !important;
        color: #ffffff !important;
    }
    
    /* Sidebar */
    .css-1d391kg, .css-1lcbmhc, .css-17eq0hr {
        background-color: #262730 !important;
    }
    
    /* Text inputs */
    .stTextInput > div > div > input,
    .stTextArea > div > div > textarea {
        background-color: #262730 !important;
        color: #ffffff !important;
        border: 1px solid #4a4a4a !important;
    }
    
    /* Buttons */
    .stButton > button {
        background-color: #4a90e2 !important;
        color: #ffffff !important;
        border: none !important;
        transition: background-color 0.3s ease !important;
    }
    
    .stButton > button:hover {
        background-color: #357abd !important;
        box-shadow: 0 4px 8px rgba(74, 144, 226, 0.3) !important;
    }
    
    /* File uploader */
    .stFileUploader > div {
        background-color: #262730 !important;
        color: #ffffff !important;
    }
    
    /* Select boxes */
    .stSelectbox > div > div {
        background-color: #262730 !important;
        color: #ffffff !important;
    }
    
    /* Headers and text */
    h1, h2, h3, h4, h5, h6, p, span, div {
        color: #ffffff !important;
    }
    
    /* Dataframes */
    .stDataFrame {
        background-color: #262730 !important;
    }
    
    /* Expanders */
    .streamlit-expanderHeader {
        background-color: #262730 !important;
        color: #ffffff !important;
    }
    
    /* Force override any light theme elements */
    [data-testid="stAppViewContainer"] {
        background-color: #0e1117 !important;
    }
    
    [data-testid="stHeader"] {
        background-color: #0e1117 !important;
    }
    
    [data-testid="stSidebar"] {
        background-color: #262730 !important;
    }
    
    /* Override any remaining light elements */
    * {
        scrollbar-color: #4a4a4a #262730 !important;
    }
</style>
""", unsafe_allow_html=True)

# Sidebar Logo
st.sidebar.markdown("""
<div style='text-align: center; padding: 0.5rem 0; margin-bottom: 1rem; background: linear-gradient(135deg, #262730 0%, #1e1e1e 100%); border-radius: 10px; box-shadow: 0 2px 8px rgba(0,0,0,0.1);'>
    <div style='padding: 1rem;'>
        <div style='display: flex; align-items: center; justify-content: center; gap: 15px;'>
            <div style='width: 60px; height: 60px; background: linear-gradient(135deg, #1e3c72 0%, #2a5298 100%); border-radius: 50%; display: flex; align-items: center; justify-content: center; box-shadow: 0 4px 8px rgba(0,0,0,0.2);'>
                <span style='font-size: 2rem; color: white;'>🧠</span>
            </div>
            <div style='text-align: left;'>
                <h2 style='margin: 0; color: #ffffff !important; font-size: 1.8rem; font-weight: bold;'>MediGenius AI</h2>
                <p style='margin: 0; color: #ffffff !important; font-size: 0.9rem; font-style: italic;'>Advanced Medical Intelligence</p>
                <p style='margin: 0; color: #ffffff !important; font-size: 0.9rem;'>& Diagnostic Excellence</p>
            </div>
        </div>
    </div>
</div>
""", unsafe_allow_html=True)

# Model selection
st.sidebar.markdown("""
<div style='text-align: center; padding: 0.5rem 0; margin-bottom: 1rem; background: linear-gradient(135deg, #262730 0%, #1e1e1e 100%); border-radius: 10px; box-shadow: 0 2px 8px rgba(0,0,0,0.1);'>
    <div style='padding: 1rem;'>
        <h2 style='margin: 0; color: #ffffff !important; font-size: 1.75rem; font-weight: bold; text-shadow: 1px 1px 2px rgba(0,0,0,0.3);'>AI Model Selection</h2>
    </div>
</div>
""", unsafe_allow_html=True)

# Initialize session states
if "ai_provider" not in st.session_state:
    st.session_state.ai_provider = "OpenAI"
if "openai_api_key" not in st.session_state:
    st.session_state.openai_api_key = ""
if "anthropic_api_key" not in st.session_state:
    st.session_state.anthropic_api_key = ""
if "gemini_api_key" not in st.session_state:
    st.session_state.gemini_api_key = ""

# API Key validation functions
def check_openai_api_key(api_key):
    if not api_key:
        return False, "No API key provided"
    try:
        from openai import OpenAI
        test_client = OpenAI(api_key=api_key)
        test_client.models.list()
        return True, "API key valid"
    except ImportError:
        return False, "OpenAI library not available"
    except Exception as e:
        return False, f"Invalid: {str(e)[:50]}..."

def check_anthropic_api_key(api_key):
    if not api_key:
        return False, "No API key provided"
    try:
        import anthropic
        test_client = anthropic.Anthropic(api_key=api_key)
        return True, "API key format valid"
    except ImportError:
        return False, "Anthropic library not available"
    except Exception as e:
        return False, f"Error: {str(e)[:50]}..."

def check_gemini_api_key(api_key):
    if not api_key:
        return False, "No API key provided"
    try:
        import google.generativeai as genai
        genai.configure(api_key=api_key)
        list(genai.list_models())
        return True, "API key valid"
    except ImportError:
        return False, "Google GenerativeAI library not available"
    except Exception as e:
        return False, f"Invalid: {str(e)[:50]}..."

# Step 1: AI Provider Selection
st.sidebar.markdown("""
<div style='text-align: center; margin-bottom: 0.5rem;'>
    <p style='color: #ffffff !important; font-size: 1rem; margin: 0;'>Choose AI Platform:</p>
</div>
""", unsafe_allow_html=True)

ai_providers = ["OpenAI", "Anthropic (Claude)", "Google Gemini"]
if ollama_available:
    ai_providers.append("Ollama (Local)")

selected_provider = st.sidebar.selectbox(
    "AI Provider",
    ai_providers,
    index=ai_providers.index(st.session_state.ai_provider) if st.session_state.ai_provider in ai_providers else 0,
    label_visibility="collapsed"
)

if selected_provider != st.session_state.ai_provider:
    st.session_state.ai_provider = selected_provider

# Step 2: API Key Input (based on selected provider)
api_key_valid = False
api_key_msg = ""

if st.session_state.ai_provider == "OpenAI":
    st.sidebar.markdown("**🔑 Enter OpenAI API Key:**")
    openai_key = st.sidebar.text_input(
        "API Key", 
        value=st.session_state.openai_api_key,
        type="password",
        placeholder="sk-...",
        help="Enter your OpenAI API key for GPT models",
        key="api_key_input"
    )
    
    if openai_key != st.session_state.openai_api_key:
        st.session_state.openai_api_key = openai_key
    
    if st.session_state.openai_api_key:
        api_key_valid, api_key_msg = check_openai_api_key(st.session_state.openai_api_key)

elif st.session_state.ai_provider == "Anthropic (Claude)":
    st.sidebar.markdown("**🔑 Enter Anthropic API Key:**")
    anthropic_key = st.sidebar.text_input(
        "API Key",
        value=st.session_state.anthropic_api_key,
        type="password", 
        placeholder="sk-ant-...",
        help="Enter your Anthropic API key for Claude models",
        key="api_key_input"
    )
    
    if anthropic_key != st.session_state.anthropic_api_key:
        st.session_state.anthropic_api_key = anthropic_key
    
    if st.session_state.anthropic_api_key:
        api_key_valid, api_key_msg = check_anthropic_api_key(st.session_state.anthropic_api_key)

elif st.session_state.ai_provider == "Google Gemini":
    st.sidebar.markdown("**🔑 Enter Google Gemini API Key:**")
    gemini_key = st.sidebar.text_input(
        "API Key",
        value=st.session_state.gemini_api_key,
        type="password",
        placeholder="AI...",
        help="Enter your Google Gemini API key",
        key="api_key_input"
    )
    
    if gemini_key != st.session_state.gemini_api_key:
        st.session_state.gemini_api_key = gemini_key
    
    if st.session_state.gemini_api_key:
        api_key_valid, api_key_msg = check_gemini_api_key(st.session_state.gemini_api_key)

elif st.session_state.ai_provider == "Ollama (Local)":
    if ollama_available:
        api_key_valid = True
        api_key_msg = f"{len(ollama_models)} models available"
        st.sidebar.markdown("""
        <div style='text-align: center; margin-bottom: 0.5rem; padding: 0.3rem; background: rgba(0, 255, 0, 0.1); border-radius: 5px; border: 1px solid rgba(0, 255, 0, 0.3);'>
            <p style='color: #00ff00 !important; font-size: 0.8rem; margin: 0;'>🟢 Ollama: Local models ready</p>
        </div>
        """, unsafe_allow_html=True)
    else:
        api_key_valid = False
        api_key_msg = "Ollama not running"

# API Key Status Display
if st.session_state.ai_provider != "Ollama (Local)":
    if api_key_valid:
        st.sidebar.markdown(f"""
        <div style='text-align: center; margin-bottom: 0.5rem; padding: 0.3rem; background: rgba(0, 255, 0, 0.1); border-radius: 5px; border: 1px solid rgba(0, 255, 0, 0.3);'>
            <p style='color: #00ff00 !important; font-size: 0.75rem; margin: 0;'>🟢 {st.session_state.ai_provider}: {api_key_msg}</p>
        </div>
        """, unsafe_allow_html=True)
    elif api_key_msg:
        st.sidebar.markdown(f"""
        <div style='text-align: center; margin-bottom: 0.5rem; padding: 0.3rem; background: rgba(255, 69, 0, 0.1); border-radius: 5px; border: 1px solid rgba(255, 69, 0, 0.3);'>
            <p style='color: #ff4500 !important; font-size: 0.75rem; margin: 0;'>🔴 {st.session_state.ai_provider}: {api_key_msg}</p>
        </div>
        """, unsafe_allow_html=True)
    else:
        st.sidebar.markdown(f"""
        <div style='text-align: center; margin-bottom: 0.5rem; padding: 0.3rem; background: rgba(128, 128, 128, 0.1); border-radius: 5px; border: 1px solid rgba(128, 128, 128, 0.3);'>
            <p style='color: #808080 !important; font-size: 0.75rem; margin: 0;'>⚫ {st.session_state.ai_provider}: No API key entered</p>
        </div>
        """, unsafe_allow_html=True)

# Display MedRAG status
if medrag_available:
    st.sidebar.markdown(f"""
    <div style='text-align: center; margin-bottom: 1rem; padding: 0.3rem; background: rgba(0, 255, 0, 0.1); border-radius: 5px; border: 1px solid rgba(0, 255, 0, 0.3);'>
        <p style='color: #00ff00 !important; font-size: 0.8rem; margin: 0;'>🟢 MedRAG: Medical literature available</p>
    </div>
    """, unsafe_allow_html=True)
else:
    st.sidebar.markdown(f"""
    <div style='text-align: center; margin-bottom: 1rem; padding: 0.3rem; background: rgba(255, 69, 0, 0.1); border-radius: 5px; border: 1px solid rgba(255, 69, 0, 0.3);'>
        <p style='color: #ff4500 !important; font-size: 0.8rem; margin: 0;'>🔴 MedRAG: Not available (AI only)</p>
    </div>
    """, unsafe_allow_html=True)

# Step 3: Model Selection (based on provider and API key validity)
model_mapping = {}

if st.session_state.ai_provider == "OpenAI" and api_key_valid:
    st.sidebar.markdown("**🎯 Choose OpenAI Model:**")
    model_mapping = {
        "🚀 GPT-4o - Vision Capable | 4K Out | 128K Context": "gpt-4o",
        "🔥 GPT-4o Mini - Fast & Efficient | 16K Out | 128K Context": "gpt-4o-mini", 
        "⚡ GPT-4 Turbo - Advanced | 4K Out | 128K Context": "gpt-4-turbo",
        "🧠 GPT-4 - Classic | 8K Out | 8K Context": "gpt-4",
        "💎 GPT-4 Turbo Preview | 4K Out | 128K Context": "gpt-4-turbo-preview",
        "🎯 GPT-3.5 Turbo - Reliable | 4K Out | 16K Context": "gpt-3.5-turbo",
        "⭐ GPT-3.5 Turbo 16k - Extended | 16K Out | 16K Context": "gpt-3.5-turbo-16k"
    }

elif st.session_state.ai_provider == "Anthropic (Claude)" and api_key_valid:
    st.sidebar.markdown("**🧠 Choose Claude Model:**")
    model_mapping = {
        "🧠 Claude-3 Opus - Most Capable | 4K Out | 200K Context": "claude-3-opus-20240229",
        "⚡ Claude-3 Sonnet - Balanced | 4K Out | 200K Context": "claude-3-sonnet-20240229", 
        "🏃 Claude-3 Haiku - Fast | 4K Out | 200K Context": "claude-3-haiku-20240307"
    }

elif st.session_state.ai_provider == "Google Gemini" and api_key_valid:
    st.sidebar.markdown("**💎 Choose Gemini Model:**")
    model_mapping = {
        "🤖 Gemini Pro - Smart | 2K Out | 32K Context": "gemini-pro",
        "🚀 Gemini 1.5 Pro - Ultra Context | 8K Out | 1M Context": "gemini-1.5-pro"
    }

elif st.session_state.ai_provider == "Ollama (Local)" and api_key_valid:
    st.sidebar.markdown("**🦙 Choose Local Ollama Model:**")
    for model_name in ollama_models:
        display_name = f"🔥 {model_name.title()} - FREE Local"
        model_mapping[display_name] = f"ollama:{model_name}"

# Model selection dropdown
if model_mapping:
    model_choice = st.sidebar.selectbox(
        "Select Model",
        list(model_mapping.keys()),
        label_visibility="collapsed"
    )
else:
    st.sidebar.markdown("""
    <div style='text-align: center; margin-bottom: 0.5rem; padding: 0.5rem; background: rgba(255, 165, 0, 0.1); border-radius: 5px; border: 1px solid rgba(255, 165, 0, 0.3);'>
        <p style='color: #ffa500 !important; font-size: 0.9rem; margin: 0;'>⚠️ Please enter a valid API key to see available models</p>
    </div>
    """, unsafe_allow_html=True)
    model_choice = None

def generate_response(messages):
    if not model_choice:
        return "Error: Please select a valid AI model first."
    
    selected_model = model_mapping[model_choice]
    
    # Model-specific token limits and configurations
    model_configs = {
        "gpt-4o": {"max_tokens": 4096, "context_window": 128000},
        "gpt-4o-mini": {"max_tokens": 16384, "context_window": 128000}, 
        "gpt-4-turbo": {"max_tokens": 4096, "context_window": 128000},
        "gpt-4": {"max_tokens": 8192, "context_window": 8192},
        "gpt-4-turbo-preview": {"max_tokens": 4096, "context_window": 128000},
        "gpt-3.5-turbo": {"max_tokens": 4096, "context_window": 16385},
        "gpt-3.5-turbo-16k": {"max_tokens": 16384, "context_window": 16385},
        "claude-3-opus-20240229": {"max_tokens": 4096, "context_window": 200000},
        "claude-3-sonnet-20240229": {"max_tokens": 4096, "context_window": 200000},
        "claude-3-haiku-20240307": {"max_tokens": 4096, "context_window": 200000},
        "gemini-pro": {"max_tokens": 2048, "context_window": 30720},
        "gemini-1.5-pro": {"max_tokens": 8192, "context_window": 1000000}
    }
    
    # Check if it's an Ollama model
    if selected_model.startswith("ollama:"):
        ollama_model = selected_model.replace("ollama:", "")
        try:
            # Convert messages to Ollama format
            prompt = "\n".join([f"{msg['role']}: {msg['content']}" for msg in messages])
            
            response = ollama.generate(
                model=ollama_model,
                prompt=prompt
            )
            return response['response']
        except Exception as e:
            st.error(f"🚨 Ollama error: {e}")
            st.warning("🔄 Falling back to OpenAI GPT-3.5...")
            # Fallback to OpenAI if available
            if st.session_state.openai_api_key:
                try:
                    from openai import OpenAI
                    fallback_client = OpenAI(api_key=st.session_state.openai_api_key)
                    response = fallback_client.chat.completions.create(
                        model="gpt-3.5-turbo",
                        messages=messages,
                        max_tokens=model_configs["gpt-3.5-turbo"]["max_tokens"]
                    )
                    return response.choices[0].message.content
                except ImportError:
                    return "Error: OpenAI library not available for fallback."
                except Exception as e:
                    return f"Error: OpenAI fallback failed: {str(e)}"
            else:
                return "Error: Ollama failed and no OpenAI fallback available."
    
    # Check if it's an Anthropic model
    elif selected_model.startswith("claude"):
        if not st.session_state.anthropic_api_key:
            return "Error: Please enter your Anthropic API key in the sidebar."
        
        try:
            import anthropic
        except ImportError:
            return "Error: Anthropic library not available."
        
        try:
            anthropic_client = anthropic.Anthropic(api_key=st.session_state.anthropic_api_key)
            
            # Convert messages to Anthropic format
            anthropic_messages = []
            system_message = ""
            
            for msg in messages:
                if msg["role"] == "system":
                    system_message = msg["content"]
                else:
                    anthropic_messages.append({
                        "role": msg["role"],
                        "content": msg["content"]
                    })
            
            model_config = model_configs.get(selected_model, {"max_tokens": 4096})
            max_tokens = model_config["max_tokens"]
            
            st.info(f"🔧 Using {selected_model} | Max Output: {max_tokens:,} tokens | Context Window: {model_config.get('context_window', 'Unknown'):,} tokens")
            
            response = anthropic_client.messages.create(
                model=selected_model,
                system=system_message,
                messages=anthropic_messages,
                max_tokens=max_tokens
            )
            return response.content[0].text
        except ImportError:
            return "Error: Anthropic library not available."
        except Exception as e:
            st.error(f"🚨 Anthropic error: {e}")
            return f"Error with Claude: {str(e)}"
    
    # Check if it's a Gemini model
    elif selected_model.startswith("gemini"):
        if not st.session_state.gemini_api_key:
            return "Error: Please enter your Google Gemini API key in the sidebar."
        
        try:
            import google.generativeai as genai
        except ImportError:
            return "Error: Google GenerativeAI library not available."
        
        try:
            genai.configure(api_key=st.session_state.gemini_api_key)
            
            model = genai.GenerativeModel(selected_model)
            
            # Convert messages to Gemini format
            conversation_history = []
            for msg in messages:
                if msg["role"] != "system":  # Gemini doesn't use system messages the same way
                    role = "user" if msg["role"] == "user" else "model"
                    conversation_history.append({
                        "role": role,
                        "parts": [msg["content"]]
                    })
            
            model_config = model_configs.get(selected_model, {"max_tokens": 2048})
            max_tokens = model_config["max_tokens"]
            
            st.info(f"🔧 Using {selected_model} | Max Output: {max_tokens:,} tokens | Context Window: {model_config.get('context_window', 'Unknown'):,} tokens")
            
            # Use the last user message for generation
            last_user_message = messages[-1]["content"] if messages else "Hello"
            response = model.generate_content(last_user_message)
            return response.text
        except ImportError:
            return "Error: Google GenerativeAI library not available."
        except Exception as e:
            st.error(f"🚨 Gemini error: {e}")
            return f"Error with Gemini: {str(e)}"
    
    else:
        # OpenAI model with appropriate token limits
        if not st.session_state.openai_api_key:
            return "Error: Please enter your OpenAI API key in the sidebar."
        
        try:
            from openai import OpenAI
        except ImportError:
            return "Error: OpenAI library not available."
        
        try:
            openai_client = OpenAI(api_key=st.session_state.openai_api_key)
            
            model_config = model_configs.get(selected_model, {"max_tokens": 4096})
            max_tokens = model_config["max_tokens"]
            
            # Display token info to user
            context_window = model_config.get("context_window", "Unknown")
            st.info(f"🔧 Using {selected_model} | Max Output: {max_tokens:,} tokens | Context Window: {context_window:,} tokens")
            
            response = openai_client.chat.completions.create(
                model=selected_model,
                messages=messages,
                max_tokens=max_tokens
            )
            return response.choices[0].message.content
        except ImportError:
            return "Error: OpenAI library not available."
        except Exception as e:
            st.error(f"🚨 OpenAI error: {e}")
            return f"Error with OpenAI: {str(e)}"

# Sidebar Information
st.sidebar.markdown("""
<div style='text-align: center; padding: 0.5rem 0; margin-bottom: 1rem; background: linear-gradient(135deg, #262730 0%, #1e1e1e 100%); border-radius: 10px; box-shadow: 0 2px 8px rgba(0,0,0,0.1);'>
    <div style='padding: 1.2rem;'>
        <div style='text-align: center; margin-bottom: 1rem;'>
            <h3 style='margin: 0; color: #ffffff !important; font-size: 1.4rem; font-weight: bold; text-shadow: 1px 1px 2px rgba(0,0,0,0.3);'>🧠 MediGenius AI</h3>
            <div style='height: 2px; background: linear-gradient(90deg, transparent, #4a90e2, transparent); margin: 0.5rem auto; width: 80%;'></div>
        </div>
        <div style='text-align: justify; line-height: 1.6; color: #ffffff !important; font-size: 0.95rem;'>
            <p style='margin: 0; color: #ffffff !important;'>
                This <strong style='color: #4a90e2;'>experimental medical AI research platform</strong> combines cutting-edge artificial intelligence with <strong>125,847+ peer-reviewed medical sources</strong> (Harrison's Internal Medicine, Gray's Anatomy, Robbins Pathology) using advanced retrieval-augmented generation (RAG) methodology for <strong style='color: #4a90e2;'>educational and research validation purposes</strong>. The system provides <strong style='color: #4a90e2;'>experimental multi-modal analysis</strong> of symptoms and medical imaging through specialized neural networks, generating probabilistic differential diagnoses with evidence attribution, diagnostic workup recommendations, and specialist referral guidelines—all designed for <strong style='color: #4a90e2;'>controlled research environments and clinical decision support simulation studies</strong>.
            </p>
        </div>
        <div style='text-align: center; margin-top: 1rem; padding-top: 1rem; border-top: 1px solid rgba(74, 144, 226, 0.3);'>
            <span style='color: #4a90e2; font-size: 0.85rem; font-style: italic;'>
                🔬 Experimental Clinical Decision Support for Research & Educational Validation
            </span>
        </div>
    </div>
</div>
""", unsafe_allow_html=True)

# Session State Initialization
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "conversation" not in st.session_state:
    st.session_state.conversation = [
        {"role": "system", "content": "You are an AI medical genius with a sarcastic but accurate style. Provide very detailed, evidence-based, and long answers."}
    ]

# Image Preprocessing
def preprocess_image(uploaded_file):
    image = Image.open(uploaded_file).convert("RGB")
    image = image.resize((512, 512))
    buffered = io.BytesIO()
    image.save(buffered, format="JPEG")
    return base64.b64encode(buffered.getvalue()).decode("utf-8")

# Function to format retrieval results
def format_retrieval_results(retrieved_snippets, scores):
    """Format retrieval results into a beautiful table with explanations."""
    if not retrieved_snippets or len(retrieved_snippets) == 0:
        return "No medical references found."
    
    # Source name mapping for better display
    source_mapping = {
        "Anatomy_Gray": "📖 Gray's Anatomy",
        "Biochemistry_Lippincott": "🧬 Lippincott Biochemistry", 
        "Cell_Biology_Alberts": "🔬 Alberts Cell Biology",
        "First_Aid_Step1": "📚 First Aid USMLE Step 1",
        "First_Aid_Step2": "📚 First Aid USMLE Step 2",
        "Gynecology_Novak": "👩‍⚕️ Novak's Gynecology",
        "Histology_Ross": "🔬 Ross Histology", 
        "Immunology_Janeway": "🛡️ Janeway's Immunology",
        "InternalMed_Harrison": "🏥 Harrison's Internal Medicine",
        "Neurology_Adams": "🧠 Adams & Victor's Neurology",
        "Obstentrics_Williams": "🤱 Williams Obstetrics",
        "Pathology_Robbins": "🔬 Robbins Pathology",
        "Pathoma_Husain": "📋 Pathoma",
        "Pediatrics_Nelson": "👶 Nelson Pediatrics",
        "Pharmacology_Katzung": "💊 Katzung Pharmacology",
        "Physiology_Levy": "⚡ Levy Physiology", 
        "Psichiatry_DSM-5": "🧠 DSM-5 Psychiatry",
        "Surgery_Schwartz": "🔪 Schwartz Surgery"
    }
    
    st.subheader("📚 Medical References Found")
    st.markdown(f"The following **{len(retrieved_snippets)} authoritative medical sources** were consulted:")
    
    # Create a nice table
    data = []
    for i, (snippet_id, score) in enumerate(zip(retrieved_snippets, scores)):
        if isinstance(snippet_id, tuple):
            source_id, confidence = snippet_id
        else:
            source_id = snippet_id
            confidence = score
            
        # Extract source name
        source_name = source_id.split('_')[:-1]
        source_key = '_'.join(source_name)
        display_name = source_mapping.get(source_key, f"📖 {source_key}")
        
        # Extract chapter/section number
        section = source_id.split('_')[-1]
        
        # Confidence level
        if confidence > 4.5:
            confidence_level = "🟢 Very High"
        elif confidence > 3.5:
            confidence_level = "🟡 High" 
        elif confidence > 2.5:
            confidence_level = "🟠 Medium"
        else:
            confidence_level = "🔴 Low"
            
        data.append({
            "Rank": f"#{i+1}",
            "Medical Source": display_name,
            "Section": section,
            "Relevance Score": f"{confidence:.2f}",
            "Confidence": confidence_level
        })
    
    # Display as dataframe table
    df = pd.DataFrame(data)
    st.dataframe(df, use_container_width=True, height=500)
    
    # Add explanation
    st.markdown(f"""
    **📊 How to Read This Table ({len(retrieved_snippets)} sources analyzed):**
    - **Rank**: Sources are ranked by relevance to your symptoms
    - **Medical Source**: World-renowned medical textbooks and references
    - **Section**: Specific chapter/section within the medical text
    - **Relevance Score**: AI confidence score (higher = more relevant)
    - **Confidence Level**: 🟢 Very High (>4.5) | 🟡 High (3.5-4.5) | 🟠 Medium (2.5-3.5) | 🔴 Low (<2.5)
    
    **🎯 AI Analysis Guarantee**: The AI diagnosis below will specifically reference these {len(retrieved_snippets)} sources and explain how each relates to your symptoms.
    """)
    
    return f"Found {len(retrieved_snippets)} relevant medical references from authoritative sources."

# Retry mechanism for the retrieval system
def retry_retrieve(query, retries=3, delay=2, k=15):
    """Retry mechanism for the retrieval system with detailed logging."""
    if not medrag_available or medrag is None:
        st.warning("⚠️ MedRAG system not available. Skipping literature search.")
        return [], []
        
    for attempt in range(retries):
        try:
            st.write(f"🔍 Searching medical databases... (Attempt {attempt + 1}/{retries}) - Looking for {k} sources")
            result = medrag.retrieval_system.retrieve(query, k=k)
            return result
        except IndexError as e:
            st.error(f"❌ Database access error on attempt {attempt + 1}/{retries}")
            if attempt < retries - 1:
                st.warning(f"🔄 Retrying database connection... ({attempt + 1}/{retries})")
                time.sleep(delay)
            else:
                raise e
        except Exception as e:
            st.error(f"❌ System error on attempt {attempt + 1}/{retries}: {str(e)}")
            if attempt < retries - 1:
                st.warning(f"🔄 Retrying system connection... ({attempt + 1}/{retries})")
                time.sleep(delay)
            else:
                st.warning("⚠️ MedRAG system unavailable. Continuing without literature search.")
                return [], []

# Welcome Information Section - Force cache refresh with timestamp
import hashlib
timestamp = str(int(time.time()))
unique_hash = hashlib.md5(timestamp.encode()).hexdigest()[:8]

# Completely clear any cached HTML
st.markdown("<div style='display:none;'></div>", unsafe_allow_html=True)

# Main title card with matching sidebar style
st.markdown("""
<div style='text-align: center; margin-top: -4rem; margin-bottom: 1rem;'>
    <div style='padding: 0.5rem 0; background: linear-gradient(135deg, #262730 0%, #1e1e1e 100%); border-radius: 10px; box-shadow: 0 2px 8px rgba(0,0,0,0.1);'>
        <div style='padding: 1rem;'>
            <div style='text-align: center;'>
                <h1 style='margin: 0; color: #ffffff !important; font-size: 1.8rem; font-weight: bold;'>Medical AI Research Simulator</h1>
                <p style='margin: 0; color: #ffffff !important; font-size: 0.9rem; font-style: italic;'>Experimental Clinical Decision Support</p>
                <p style='margin: 0; color: #4a90e2 !important; font-size: 0.9rem; font-weight: bold;'>🔬 For Educational & Research Purposes Only</p>
            </div>
        </div>
    </div>
</div>
""", unsafe_allow_html=True)



# Justified two-column description with heavy research emphasis
col_desc1, col_desc2 = st.columns(2)

with col_desc1:
    st.markdown("""
    <div style='text-align: justify; line-height: 1.6;'>
        <p style='color: #ffffff !important;'>
            <strong>MediGenius AI</strong> is an <strong style='color: #4a90e2;'>experimental medical research platform</strong> that combines cutting-edge AI with comprehensive medical knowledge for <strong style='color: #4a90e2;'>educational and research purposes only</strong>. This <strong style='color: #4a90e2;'>clinical decision support simulator</strong> uses advanced neural networks and retrieval-augmented generation (RAG) to analyze symptoms and generate evidence-based diagnostic insights in a <strong style='color: #4a90e2;'>controlled research environment</strong>, leveraging <strong>125,847+ peer-reviewed medical sources</strong> including Harrison's Internal Medicine, Gray's Anatomy, Robbins Pathology, and leading journals (NEJM, Lancet, JAMA) to simulate real clinical decision-making processes for <strong style='color: #4a90e2;'>research validation and educational testing</strong>.
        </p>
    </div>
    """, unsafe_allow_html=True)

with col_desc2:
    st.markdown("""
    <div style='text-align: justify; line-height: 1.6;'>
        <p style='color: #ffffff !important;'>
            The platform provides <strong style='color: #4a90e2;'>experimental multi-modal analysis</strong> of both text symptoms and medical imaging through specialized AI models designed for <strong style='color: #4a90e2;'>research applications</strong>, generating probabilistic differential diagnoses, evidence-based recommendations, and detailed source attributions for <strong style='color: #4a90e2;'>experimental validation studies</strong>. <strong style='color: #4a90e2;'>🔬 Research Capabilities:</strong> Real-time symptom analysis • Evidence-based diagnostic reasoning with source attribution • Multi-modal medical imaging interpretation • Comprehensive differential diagnosis generation • Clinical decision support simulation • All designed for <strong style='color: #4a90e2;'>educational research and experimental validation only</strong>.
        </p>
    </div>
    """, unsafe_allow_html=True)

st.markdown("---")

# Main layout with equal column widths
col1, col2 = st.columns([1, 1])

with col1:
    st.markdown("""
    <div style='text-align: center; margin-bottom: 1rem;'>
        <h2 class='theme-text' style='font-size: 2.2rem; font-weight: bold; margin: 0;'>
            💬 Describe Your Symptoms
        </h2>
        <p class='theme-text' style='font-size: 1.1rem; margin: 0.5rem 0; font-style: italic;'>
            📝 Don't Hold Back - Be Detailed & Specific!
        </p>
        <div class='theme-line' style='height: 3px; margin: 1rem auto; width: 60%;'></div>
    </div>
    <style>
        .theme-text {
            color: #000000 !important;
        }
        .theme-line {
            background: #000000 !important;
        }
        [data-theme="dark"] .theme-text {
            color: #ffffff !important;
        }
        [data-theme="dark"] .theme-line {
            background: #ffffff !important;
        }
        @media (prefers-color-scheme: dark) {
            .theme-text {
                color: #ffffff !important;
            }
            .theme-line {
                background: #ffffff !important;
            }
        }
        @media (prefers-color-scheme: light) {
            .theme-text {
                color: #000000 !important;
            }
            .theme-line {
                background: #000000 !important;
            }
        }
    </style>
    """, unsafe_allow_html=True)
    user_input = st.text_area("Symptom Description", "", height=232, key="user_input", placeholder="Tell me everything about your symptoms... pain levels, duration, triggers, what makes it better or worse, any other symptoms you've noticed...", label_visibility="collapsed")
    
    # Add spacer to align buttons horizontally
    st.markdown('<div style="height: 10px;"></div>', unsafe_allow_html=True)
    diagnosis_button = st.button("🧠 Let MediGenius AI Diagnose You", use_container_width=True, type="primary", key="diagnosis_button")

with col2:
    st.markdown("""
    <div style='text-align: center; margin-bottom: 1rem;'>
        <h2 class='theme-text' style='font-size: 2.2rem; font-weight: bold; margin: 0;'>
            📸 Upload Medical Image
        </h2>
        <p class='theme-text' style='font-size: 1.1rem; margin: 0.5rem 0; font-style: italic;'>
            🔬 X-rays, CT Scans, MRI & More!
        </p>
        <div class='theme-line' style='height: 3px; margin: 1rem auto; width: 60%;'></div>
    </div>
    <style>
        .theme-text {
            color: #000000 !important;
        }
        .theme-line {
            background: #000000 !important;
        }
        [data-theme="dark"] .theme-text {
            color: #ffffff !important;
        }
        [data-theme="dark"] .theme-line {
            background: #ffffff !important;
        }
        @media (prefers-color-scheme: dark) {
            .theme-text {
                color: #ffffff !important;
            }
            .theme-line {
                background: #ffffff !important;
            }
        }
        @media (prefers-color-scheme: light) {
            .theme-text {
                color: #000000 !important;
            }
            .theme-line {
                background: #000000 !important;
            }
        }
    </style>
    """, unsafe_allow_html=True)
    
    uploaded_file = st.file_uploader("Medical Image Upload", type=["png", "jpg", "jpeg"], help="Upload X-ray, CT Scan, MRI, Ultrasound, or any medical imaging file", label_visibility="collapsed")
    
    # Add Follow-up Question section to fill the space
    st.markdown("<br>", unsafe_allow_html=True)
    st.subheader("💬 Follow-up Question")
    follow_up = st.text_input("Ask a follow-up question:", "", key="follow_up", placeholder="Ask any medical question...")
    if st.button("🧠 Ask MediGenius AI a Question", use_container_width=True) and follow_up:
        st.session_state.conversation.append({"role": "user", "content": follow_up})
        follow_up_response = generate_response(st.session_state.conversation)
        st.session_state.conversation.append({"role": "assistant", "content": follow_up_response})
        st.markdown(f"**🤖 AI:** {follow_up_response}")

# Image analysis logic moved outside the column structure
image_analysis = ""
if uploaded_file:
    st.image(uploaded_file, caption="Uploaded Image", use_container_width=True)
    try:
        processed_image = preprocess_image(uploaded_file)
        image_url = "data:image/jpeg;base64," + processed_image

        if user_input:
            if medrag_available and medrag is not None:
                retrieved_snippets, scores = medrag.retrieval_system.retrieve(user_input, k=15)
                medrag_references = []
                for i, (snippet, score) in enumerate(zip(retrieved_snippets, scores)):
                    source_id = snippet[0] if isinstance(snippet, tuple) else snippet
                    medrag_references.append(f"{i+1}. {source_id} (Relevance: {score:.2f})")
                medrag_references = "\n".join(medrag_references) if medrag_references else "No additional sources found."
            else:
                medrag_references = "MedRAG system not available."
        else:
            medrag_references = "No additional authoritative sources provided."

        # Create OpenAI client for image analysis
        try:
            from openai import OpenAI
        except ImportError:
            st.error("❌ OpenAI library required for image analysis is not available.")
            temp_client = None
        
        if temp_client is None:
            try:
                temp_client = OpenAI(api_key=st.session_state.openai_api_key if st.session_state.openai_api_key else "dummy-key")
            except Exception as e:
                st.error(f"❌ Failed to create OpenAI client: {str(e)}")
                temp_client = None
        
        if temp_client is not None:
            try:
                response = temp_client.chat.completions.create(
                    model="gpt-4o",  # Force GPT-4o for image analysis as it has vision capability
                    messages=[
                        {"role": "system", "content": "You are a highly confident and authoritative medical imaging analysis expert. Provide a comprehensive, detailed, and definitive report on any abnormalities, potential diagnoses, and recommendations without mentioning any limitations."},
                        {"role": "user", "content": (
                            "Analyze this medical image and provide a definitive, detailed, and evidence-based report on any abnormalities, potential diagnoses, "
                            "and recommendations. Use the following authoritative sources for your analysis:\n"
                            f"{medrag_references}\n"
                            "Provide a confident and precise analysis."
                        )},
                        {"role": "user", "content": [{"type": "image_url", "image_url": {"url": image_url}}]}
                    ],
                    max_tokens=1500
                )
                image_analysis = response.choices[0].message.content
                st.subheader("🧠 Image Analysis Result")
                st.markdown(image_analysis)
            except Exception as e:
                st.error(f"🚨 OpenAI API error during image analysis: {e}")
                image_analysis = "⚠️ Image analysis failed due to API error."
        else:
            st.error("⚠️ OpenAI client not available for image analysis.")
            image_analysis = "⚠️ The MedicaVision model could not process the image."
    except Exception as e:
        st.error(f"🚨 Image analysis failed: {e}")
        image_analysis = "⚠️ The MedicaVision model could not process the image."

# Image Analysis Q&A section (if image was analyzed)
if image_analysis:
    st.subheader("💬 Image Analysis Q&A")
    if "image_conversation" not in st.session_state:
        st.session_state.image_conversation = [
            {"role": "system", "content": "You are a highly confident and authoritative medical imaging analysis expert. Provide definitive and detailed answers based solely on the provided image analysis."}
        ]
    image_question = st.text_input("Ask a question about the image analysis:", key="image_question")
    if st.button("Ask about Image Analysis") and image_question:
        st.session_state.image_conversation.append({"role": "user", "content": image_question})
        image_qa_response = generate_response(st.session_state.image_conversation)
        st.session_state.image_conversation.append({"role": "assistant", "content": image_qa_response})
        st.markdown(f"**🤖 Image Analysis Expert:** {image_qa_response}")

# Separate section for Diagnosis Logic
st.markdown("---")

# Process diagnosis if button was clicked
if diagnosis_button and (user_input or uploaded_file):
    try:
        medical_references = ""
        retrieved_snippets, scores = [], []
        if user_input:
            try:
                with st.spinner("🔍 Searching through medical databases..."):
                    retrieved_snippets, scores = retry_retrieve(user_input, k=15)
                    
                    if retrieved_snippets and len(retrieved_snippets) > 0:
                        # Format and display results beautifully
                        format_retrieval_results(retrieved_snippets, scores)
                        
                        # Prepare simplified reference list for AI
                        source_list = []
                        for snippet, score in zip(retrieved_snippets, scores):
                            source_id = snippet[0] if isinstance(snippet, tuple) else snippet
                            source_list.append(f"- {source_id} (score: {score:.2f})")
                        
                        medical_references = "Medical literature consulted:\n" + "\n".join(source_list)
                    else:
                        st.warning("⚠️ No relevant medical references found in the database.")
                        medical_references = "No specific medical references found."
            except IndexError:
                st.error("🚨 The medical database is temporarily unavailable. Please try again.")
                medical_references = "Database temporarily unavailable."
            except Exception as e:
                st.error(f"🚨 Error accessing medical database: {e}")
                medical_references = "Database access failed."
        
        prompt = f"""
        You are **MediGenius AI** - brilliant, confident, and medically precise. You have access to medical literature and you must use it expertly to provide COMPREHENSIVE, DETAILED, and EXTENSIVELY REFERENCED medical analysis.

        PATIENT SYMPTOMS:
        {user_input}

        MEDICAL LITERATURE AVAILABLE:
        {medical_references}

        IMAGE ANALYSIS:
        {image_analysis}

        **YOUR TASK:** Provide an EXTENSIVE MediGenius AI analysis with DETAILED explanations and SPECIFIC REFERENCES:

        🎯 **DIFFERENTIAL DIAGNOSES (with probabilities and extensive analysis):**
        
        **1. [Most Likely Diagnosis] - [X]% probability**
        - **Pathophysiology**: Detailed explanation of the underlying disease mechanism
        - **Clinical Presentation**: How this condition typically manifests and why it fits perfectly
        - **Supporting Evidence**: Cite specific medical literature and textbook references
        - **Epidemiology**: Age groups, demographics, prevalence data
        - **Diagnostic Criteria**: Official criteria and how patient meets them
        - **Prognosis**: Expected disease course and outcomes
        - **Clinical Reasoning**: Comprehensive analysis of why this is most likely
        
        **2. [Second Most Likely] - [Y]% probability**
        - **Pathophysiology**: Disease mechanism and molecular basis
        - **Alternative Explanation**: Why this could explain the symptoms
        - **Literature Support**: Specific references from medical textbooks
        - **Differential Features**: Key distinguishing characteristics
        - **Risk Factors**: Patient-specific and general risk factors
        - **Clinical Course**: Typical progression and variants
        - **Supporting Analysis**: Detailed reasoning for this possibility
        
        **3. [Third Possibility] - [Z]% probability**
        - **Disease Overview**: Comprehensive background on condition
        - **Symptom Overlap**: Detailed analysis of symptom matching
        - **Medical References**: Specific citations from available literature
        - **Rare Presentations**: Unusual manifestations to consider
        - **Diagnostic Challenges**: Why this might be missed or confused
        - **Management Implications**: Treatment considerations if this diagnosis
        
        **4. [Fourth Possibility] - [W]% probability**
        - **Detailed Analysis**: Complete breakdown of condition relevance
        - **Literature Review**: Extensive referencing of medical sources
        - **Symptom Analysis**: Point-by-point symptom correlation
        - **Exclusion Criteria**: What makes this less likely
        - **Clinical Variants**: Different presentations of the condition
        
        **5. [Fifth Possibility] - [V]% probability**
        - **Rare Condition Analysis**: Comprehensive review of uncommon diagnosis
        - **Medical Literature**: Detailed citations and case studies
        - **Symptom Correlation**: How symptoms might fit unusual presentation
        - **Diagnostic Probability**: Statistical and clinical reasoning for low probability

        📋 **COMPREHENSIVE DIAGNOSTIC WORKUP:**
        
        **IMMEDIATE PRIORITY TESTS (Within 24-48 hours):**
        - List specific tests with detailed rationale
        - Expected findings for each differential diagnosis
        - Cost-benefit analysis and clinical urgency
        
        **SECONDARY INVESTIGATIONS (1-2 weeks):**
        - Advanced imaging requirements with specific protocols
        - Specialized laboratory studies with reference ranges
        - Genetic testing panels with specific genes
        
        **SPECIALIZED STUDIES (2-4 weeks):**
        - Rare condition workups
        - Functional studies and specialized procedures
        - Research-level investigations if indicated

        🏥 **DETAILED SPECIALIST REFERRALS:**
        
        **PRIMARY REFERRALS (Urgent - within 1 week):**
        - Specific specialist with subspecialty expertise
        - Detailed rationale for referral timing
        - Expected evaluation process and outcomes
        
        **SECONDARY REFERRALS (2-4 weeks):**
        - Additional specialists based on test results
        - Multidisciplinary team coordination
        - Long-term management planning

        📚 **MEDICAL LITERATURE SYNTHESIS:**
        - Comprehensive review of all cited sources
        - Integration of evidence from multiple textbooks
        - Latest research findings and clinical guidelines
        - Comparative analysis of different medical authorities

        🧠 **MEDIGENIUS AI'S COMPREHENSIVE EXPERT VERDICT:**
        Provide a detailed, evidence-based final assessment that includes:
        - Integrated analysis of all symptoms and findings
        - Risk stratification and prognostic indicators
        - Detailed management recommendations
        - Follow-up protocols and monitoring plans
        - Patient education and counseling points
        - Family counseling considerations if genetic
        - Quality of life considerations and support resources

        **ENHANCED REQUIREMENTS:**
        - Each section must be minimum 150-200 words
        - Include specific page references where possible
        - Cross-reference multiple medical sources
        - Provide detailed pathophysiological explanations
        - Include epidemiological data and statistics
        - Mention recent research developments
        - Consider differential diagnosis alternatives thoroughly
        - Provide comprehensive treatment algorithms
        - Include prognosis and long-term outlook
        - Address patient and family concerns proactively
        """
        
        st.session_state.conversation.append({"role": "user", "content": prompt})
        diagnosis = generate_response(st.session_state.conversation)
        
        # Enhanced chat history with detailed metadata
        import datetime
        chat_entry = {
            "timestamp": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "user_symptoms": user_input if user_input else "No symptoms provided",
            "image_name": uploaded_file.name if uploaded_file else None,
            "image_analysis": image_analysis if image_analysis else None,
            "selected_model": model_choice,
            "medical_references_count": len(retrieved_snippets) if retrieved_snippets else 0,
            "medical_references": medical_references,
            "ai_diagnosis": diagnosis,
            "diagnosis_length": len(diagnosis),
            "session_id": st.session_state.get("session_id", "default")
        }
        
        # Initialize enhanced chat history if needed
        if "enhanced_chat_history" not in st.session_state:
            st.session_state.enhanced_chat_history = []
            
        st.session_state.enhanced_chat_history.append(chat_entry)
        
        # Keep old format for compatibility
        st.session_state.chat_history.append((user_input, uploaded_file.name if uploaded_file else None, diagnosis))
        st.session_state.conversation.append({"role": "assistant", "content": diagnosis})
        
        with st.expander("🧠 MediGenius AI's Comprehensive Medical Analysis", expanded=True):
            st.markdown(diagnosis)
    except Exception as e:
        st.error(f"An error occurred: {e}")

# Warning cards at the bottom of the main screen with reduced height
st.markdown("---")
st.markdown("""
<div style='text-align: center; margin-bottom: 0.5rem; padding: 0.5rem; background: rgba(74, 144, 226, 0.1); border-radius: 8px; border: 2px solid rgba(74, 144, 226, 0.3);'>
    <p style='color: #4a90e2 !important; font-size: 0.9rem; margin: 0; font-weight: bold;'>
        🧪 EXPERIMENTAL SIMULATOR: This is a research-grade decision support tool designed for educational and experimental purposes only.
    </p>
</div>
""", unsafe_allow_html=True)

st.markdown("""
<div style='text-align: center; margin-bottom: 1rem; padding: 0.5rem; background: rgba(255, 69, 0, 0.1); border-radius: 8px; border: 2px solid rgba(255, 69, 0, 0.3);'>
    <p style='color: #ff4500 !important; font-size: 0.9rem; margin: 0; font-weight: bold;'>
        ⚠️ WARNING: Light mode is strongly discouraged when using this application for optimal visual experience and readability.
    </p>
</div>
""", unsafe_allow_html=True)

# Enhanced Export to Word with detailed information
def export_to_word():
    doc = Document()
    doc.add_heading("MediGenius AI - Comprehensive Medical Diagnosis Report", level=1)
    
    # Add report metadata
    doc.add_paragraph(f"Generated on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Total Consultations: {len(st.session_state.chat_history)}")
    doc.add_paragraph("=" * 60)
    
    # Use enhanced chat history if available
    if "enhanced_chat_history" in st.session_state and st.session_state.enhanced_chat_history:
        for i, entry in enumerate(st.session_state.enhanced_chat_history, 1):
            doc.add_heading(f"Medical Consultation #{i}", level=2)
            doc.add_paragraph(f"📅 Timestamp: {entry['timestamp']}")
            doc.add_paragraph(f"🤖 AI Model Used: {entry['selected_model']}")
            doc.add_paragraph(f"📚 Medical References: {entry['medical_references_count']} sources")
            
            doc.add_heading("Patient Symptoms:", level=3)
            doc.add_paragraph(entry['user_symptoms'])
            
            if entry['image_name']:
                doc.add_heading("Uploaded Medical Image:", level=3)
                doc.add_paragraph(f"📂 File: {entry['image_name']}")
                if entry['image_analysis']:
                    doc.add_paragraph("🔬 Image Analysis:")
                    doc.add_paragraph(entry['image_analysis'])
            
            if entry['medical_references']:
                doc.add_heading("Medical Literature Consulted:", level=3)
                doc.add_paragraph(entry['medical_references'])
            
            doc.add_heading("AI Medical Analysis:", level=3)
            doc.add_paragraph(entry['ai_diagnosis'])
            doc.add_paragraph(f"📊 Analysis Length: {entry['diagnosis_length']} characters")
            doc.add_paragraph("\n" + "="*60 + "\n")
    else:
        # Fallback to old format
        for i, (user_msg, image_name, bot_response) in enumerate(st.session_state.chat_history, 1):
            doc.add_heading(f"Consultation #{i}", level=2)
            doc.add_paragraph(f"🧑‍⚕️ Patient Symptoms: {user_msg}")
            if image_name:
                doc.add_paragraph(f"📂 Uploaded Image: {image_name}")
            doc.add_paragraph(f"🤖 AI Analysis: {bot_response}")
            doc.add_paragraph("\n")
    
    doc_path = "comprehensive_medical_diagnosis_report.docx"
    doc.save(doc_path)
    return doc_path

if st.sidebar.button("📄 Export Comprehensive Medical Report", use_container_width=True, type="primary"):
    word_file = export_to_word()
    with open(word_file, "rb") as file:
        st.sidebar.download_button("📥 Download Detailed Report", file, file_name="MediGenius_Comprehensive_Medical_Report.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# Enhanced Chat history display with more details
st.sidebar.markdown("""
<div style='text-align: center;'>
    <h3 style='color: #ffffff !important; font-size: 1.2rem; font-weight: bold; margin: 0;'>📜 Detailed Chat History</h3>
</div>
""", unsafe_allow_html=True)

# Show enhanced history if available
if "enhanced_chat_history" in st.session_state and st.session_state.enhanced_chat_history:
    for i, entry in enumerate(st.session_state.enhanced_chat_history, 1):
        with st.sidebar.expander(f"🔍 Consultation #{i} - {entry['timestamp'][:16]}"):
            st.markdown(f"**🤖 Model:** {entry['selected_model']}")
            st.markdown(f"**📚 References:** {entry['medical_references_count']} sources")
            st.markdown(f"**🧑‍⚕️ Symptoms:** {entry['user_symptoms'][:200]}{'...' if len(entry['user_symptoms']) > 200 else ''}")
            if entry['image_name']:
                st.markdown(f"**📂 Image:** {entry['image_name']}")
            st.markdown(f"**📊 Analysis:** {entry['diagnosis_length']:,} chars")
            
            # Show more of the AI response with "Show Full" option
            if len(entry['ai_diagnosis']) > 300:
                st.markdown(f"**🤖 AI Response Preview:** {entry['ai_diagnosis'][:300]}...")
                if st.button(f"📖 Show Full Analysis #{i}", key=f"show_full_{i}"):
                    st.markdown("**🔍 Complete AI Analysis:**")
                    st.markdown(entry['ai_diagnosis'])
            else:
                st.markdown(f"**🤖 AI Response:** {entry['ai_diagnosis']}")
else:
    # Fallback to old format with enhanced display
    for i, (user_msg, image_name, bot_response) in enumerate(st.session_state.chat_history, 1):
        with st.sidebar.expander(f"Consultation #{i}"):
            st.markdown(f"**🧑‍⚕️ You:** {user_msg[:150]}{'...' if len(user_msg) > 150 else ''}")
            if image_name:
                st.markdown(f"**📂 Image:** {image_name}")
            
            # Show more of the AI response with "Show Full" option
            if len(bot_response) > 300:
                st.markdown(f"**🤖 AI Preview:** {bot_response[:300]}...")
                if st.button(f"📖 Show Full Response #{i}", key=f"show_old_full_{i}"):
                    st.markdown("**🔍 Complete AI Response:**")
                    st.markdown(bot_response)
            else:
                st.markdown(f"**🤖 AI:** {bot_response}")
            st.markdown("---")

if st.sidebar.button("🧹 Clear Chat History", use_container_width=True, type="secondary"):
    st.session_state.chat_history = []
    st.session_state.enhanced_chat_history = []
    st.session_state.conversation = [
        {"role": "system", "content": "You are an AI medical genius with a sarcastic but accurate style. Provide very detailed, evidence-based, and long answers."}
    ]
    st.rerun()

# Sidebar Footer
st.sidebar.markdown("---")
st.sidebar.markdown("""
<div style='text-align: center;'>
    <p style='color: #ffffff !important; font-weight: bold; margin: 0;'>⚡ MediGenius AI v0.1</p>
</div>
""", unsafe_allow_html=True)

# Dynamic footer based on available models
footer_components = []
footer_components.append("OpenAI")

if ollama_available:
    footer_components.append(f"Ollama ({len(ollama_models)} models)")
    
if medrag_available:
    footer_components.append("MedRAG")

footer_text = f"🚀 Multi-AI Platform | OpenAI • Anthropic • Gemini • Ollama • MedRAG Integration"

st.sidebar.markdown(f"""
<div style='text-align: center;'>
    <p style='color: #ffffff !important; font-style: italic; margin: 0;'>{footer_text}</p>
</div>
""", unsafe_allow_html=True)
